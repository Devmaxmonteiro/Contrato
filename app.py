from openpyxl import load_workbook
from docx import Document
from datetime import datetime

# - Passar as informações da planilha para o arquivo word
planilha_fornecedores = load_workbook('./fornecedores.xlsx')
pagina_fornecedores = planilha_fornecedores['Sheet1']

for linha in pagina_fornecedores.iter_rows(min_row=2, values_only=True):
    nome_empresa, endereco, cidade, estado, cep, telefone, email, setor = linha  # unpacking

    arquivo_word = Document()
    arquivo_word.add_heading('Contrato de Prestação de serviço', 0)

    data_atual = datetime.now().strftime("%d/%m/%Y")

    texto_contrato = f"""
    Este contrato de prestação de serviços é feito entre {nome_empresa}, com endereço em {endereco}, 
    {cidade}, {estado}, CEP {cep}, doravante denominado FORNECEDOR, e a empresa CONTRATANTE.

    Pelo presente instrumento particular, as partes têm, entre si, justo e acordado o seguinte:

    1. OBJETO DO CONTRATO
    O FORNECEDOR compromete-se a fornecer à CONTRATANTE os serviços/material de acordo com as especificações acordadas, respeitando os padrões de qualidade e os prazos estipulados.

    2. PRAZO
    Este contrato tem prazo de vigência de 12 (doze) meses, iniciando-se na data de sua assinatura, podendo ser renovado conforme acordo entre as partes.

    3. VALOR E FORMA DE PAGAMENTO
    O valor dos serviços prestados será acordado conforme as demandas da CONTRATANTE e a capacidade de entrega do FORNECEDOR. Os pagamentos serão realizados mensalmente, mediante apresentação de nota fiscal.

    4. CONFIDENCIALIDADE
    Todas as informações trocadas entre as partes durante a vigência deste contrato serão tratadas como confidenciais.

    Para firmeza e como prova de assim haverem justo e contratado, as partes assinam o presente contrato em duas vias de igual teor e forma.

    FORNECEDOR: {nome_empresa}
    E-mail: {email}

    CONTRATANTE: [NOME CONTRATANTE]
    E-mail: [E-MAIL CONTRATANTE]

    {cidade}, {data_atual}
    """

    arquivo_word.add_paragraph(texto_contrato)
    
    # Salvando o documento
    arquivo_word.save(f'contrato_{nome_empresa}.docx')


